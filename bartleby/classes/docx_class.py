import bartleby.configuration as conf
import google.auth

from docx import Document
from docx.shared import Pt
from googleapiclient.discovery import build
from googleapiclient.errors import HttpError
from googleapiclient.http import MediaFileUpload

class Docx:
    '''Class to hold objects related to document output'''

    def __init__(self):

        # Document generation stuff.
        self.document_path = conf.DOCUMENTS_PATH
        self.template = Document(f'{self.document_path}/{conf.docx_template_file}')

        # Google API service account credentials set via venv
        self.service_account_credentials, _ = google.auth.default()

    def generate(self, user, first_message_number, last_message_number):
        '''Recovers text from chat, formats as docx and
        pushes to google drive'''

        # Set gdrive folder id from user
        self.gdrive_folder_id = user.gdrive_folder_id

        # Load docx template
        self.load_template()

        # Add heading 
        result = self.template.add_paragraph(user.document_title, style = 'Heading 1')

        # Decide how to select message from the chat history

        # If there is no second message range parameter, just get
        # the message specified by the first message range parameter
        if last_message_number == None:
            body = user.messages[-first_message_number]['content']

            # Split on newline so we can format paragraphs correctly
            paragraphs = body.split('\n')

        else:
            # Empty holder for paragraphs
            paragraphs = []
            body = []

            # Loop on messages in range
            while last_message_number >= first_message_number:
                #for message in user.messages[-last_message_number:-first_message_number]:
                
                # Get message content
                content = user.messages[-last_message_number]['content']
                last_message_number -= 1

                # Spit on any newlines and add to paragraphs list 
                paragraphs.extend(content.split('\n'))
                body.append(content)

            body = ' '.join(body)

        paragraph_count = 0

        # Loop on paragraphs and add them to doc
        for paragraph in paragraphs:

            # Make sure this 'paragraph' has content, i.e. it wasn't
            # the result of splitting a multiple newline
            if len(paragraph) > 0:

                # Add the paragraph to the document
                paragraph_count += 1
                result = self.template.add_paragraph(paragraph)

                # Deal with spacing between paragraphs. If the body contains
                # multiple paragraphs and this is not the last one, add some
                # space after it.
                if (len(paragraphs) > 1) and (paragraph_count < len(body)):
                    result.paragraph_format.space_after = Pt(6)

        # Format output file name for docx file
        output_filename = f'{user.document_title.replace(" ", "_")}.docx'

        # Save the docx
        self.template.save(f'{self.document_path}/{output_filename}')

        # Upload the docx to gdrive
        result = self.upload(output_filename)

    def load_template(self):
        '''Loads and returns empty docx document
        with some pre-defined styles'''

        # Empty out template
        lines = self.template.paragraphs

        for line in lines:
            p = line._element
            p.getparent().remove(p)
            p._p = p._element = None

    def upload(self, filename):
        '''Takes file name as string. uploads to google drive
        Returns : Id's of the file uploaded'''

        try:
            # create drive api client
            service = build(
                'drive', 
                'v3', 
                credentials = self.service_account_credentials
            )

            # Set file metadata
            file_metadata = {
                'name': filename,
                'parents': [f'{self.gdrive_folder_id}']
            }

            # Create media
            file_type = 'text/plain'

            if 'docx' in filename:
                file_type = 'application/vnd.openxmlformats-officedocument.wordprocessingml.document'

            media = MediaFileUpload(
                f'{conf.DOCUMENTS_PATH}/{filename}',
                mimetype = file_type
            )
            
            # Do the upload
            file = service.files().create(
                body = file_metadata, 
                media_body = media,
                fields = 'id'
            ).execute()

        # Catch http errors
        except HttpError as error:
            print(f'An error occurred while uploading file: {error}')
            file = None

        # Done
        return file.get('id')