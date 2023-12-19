import bartleby.configuration as conf
import google.auth

from docx import Document
from docx.shared import Pt
from googleapiclient.discovery import build
from googleapiclient.errors import HttpError
from googleapiclient.http import MediaFileUpload

class Docx:
    '''Class to hold objects related to document output'''

    def __init__(self):

        # Document generation stuff.
        self.title = conf.default_title
        self.document_output_path = conf.DOCUMENTS_PATH
        self.docx_template_file = conf.docx_template_file
        self.gdrive_folder_id = conf.gdrive_folder_id

    async def generate(self, llm_instance):
        '''Recovers bot generated text from chat, formats as docx and
        pushes to google drive'''

        # Load docx template
        self.load_template()

        # Add heading 
        result = self.template.add_paragraph(self.title, style = 'Heading 1')

        # Get last message in chain
        body = llm_instance.messages[-1]['content']

        # Split on newline so we can format paragraphs correctly
        paragraphs = body.split('\n')

        paragraph_count = 0

        for paragraph in paragraphs:

            # Make sure this 'paragraph' has content, i.e. it wasn't
            # the result of splitting a multiple newline
            if len(paragraph) > 0:

                # Add the paragraph to the document
                paragraph_count += 1
                result = self.template.add_paragraph(paragraph)

                # Deal with spacing between paragraphs. If the body contains
                # multiple paragraphs and this is not the last one, add some
                # space after it.
                if (len(paragraphs) > 1) and (paragraph_count < len(body)):
                    result.paragraph_format.space_after = Pt(6)

        # Format output file name for docx file
        output_filename = f'{self.title.replace(" ", "_")}.docx'

        # Save the docx
        self.template.save(f'{self.document_output_path}/{output_filename}')

        # Upload the docx to gdrive
        result = self.upload(output_filename)


    def load_template(self):
        '''Loads and returns empty docx document
        with some pre-defined styles'''

        # Load docx template
        self.template = Document(f'{conf.DATA_PATH}/{self.docx_template_file}')

        # Empty out template
        lines = self.template.paragraphs

        for line in lines:
            p = line._element
            p.getparent().remove(p)
            p._p = p._element = None

    def upload(self, filename):
        '''Takes file name as string. uploads to google drive
        Returns : Id's of the file uploaded'''

        # Uses google API service account credentials set via venv
        service_account_credentials, result = google.auth.default()

        try:
            # create drive api client
            service = build(
                'drive', 
                'v3', 
                credentials = service_account_credentials
            )

            # Set file metadata
            file_metadata = {
                'name': filename,
                'parents': [f'{self.gdrive_folder_id}']
            }

            # Create media
            file_type = 'text/plain'

            if 'docx' in filename:
                file_type = 'application/vnd.openxmlformats-officedocument.wordprocessingml.document'

            media = MediaFileUpload(
                f'{conf.DOCUMENTS_PATH}/{filename}',
                mimetype = file_type
            )
            
            # Do the upload
            file = service.files().create(
                body = file_metadata, 
                media_body = media,
                fields = 'id'
            ).execute()

        # Catch http errors
        except HttpError as error:
            print(f'An error occurred while uploading file: {error}')
            file = None

        # Done
        return file.get('id')